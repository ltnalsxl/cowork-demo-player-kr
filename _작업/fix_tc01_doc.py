# -*- coding: utf-8 -*-
"""TC-01 산출물 Word 문서의 대체어를 사내 교육팀 표현으로 바꾼다.

익명화 1차에서 "금융 고객 C" 같은 대체어를 넣었는데, 그러면 고객사에 교육을
파는 회사로 읽힌다. 사내 교육팀 한 팀의 일로 옮긴다.

run 단위로 바꾼다. 문단이나 셀 통째로 텍스트를 갈아 끼우면 서식이 날아간다.
치환 문자열이 run 경계에 걸쳐 있으면 그 문단의 run을 합쳐 처리한다.
"""
import re
import sys
import os

from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))
DOC = os.path.join(BASE, '..', 'assets', 'artifacts', 'tc01', '주간업무보고.docx')

# 긴 것부터 둔다. 짧은 것이 먼저 걸리면 긴 표현이 반쪽만 바뀐다.
MAP = [
    ('금융 고객 C 전사 오픈 교육', '전사 신규 시스템 기초 교육'),
    ('금융 고객 C – 전사 오픈 기초 교육', '영업본부 – 전사 신규 시스템 기초 교육'),
    ('금융 고객 C', '영업본부'),
    ('유통 고객 B – 협업 도구 활용 교육', '생산본부 – 협업 도구 활용 교육'),
    ('유통 고객 B 실습 세션', '생산본부 실습 세션'),
    ('유통 고객 B', '생산본부'),
    ('서비스 고객 D – 도입 프로그램', 'IT본부 – 학습관리시스템 도입'),
    ('서비스 고객 D', 'IT본부'),
    ('제조 고객 A – 지방 사업장 다일차 교육', '지방 사업장 – 다일차 실무 교육'),
    ('제조 고객 A 워크숍', '지방 사업장 워크숍'),
    ('제조 고객 A 지방 사업장 다일차 교육', '지방 사업장 다일차 교육'),
    ('제조 고객 A', '지방 사업장'),
    ('보험 고객 E – 도입 세션', '신임 팀장 과정 – 도입 세션'),
    ('보험 고객 E', '신임 팀장 과정'),
    ('고객 지원 요청 대응 (통신 고객 F, 커머스 고객 G)', '교육 문의 대응 (재무본부, 인사본부)'),
    ('통신 고객 F', '재무본부'),
    ('커머스 고객 G', '인사본부'),
    ('도입 프로그램이 서비스 고객 D에서 승인', '학습관리시스템이 IT본부에서 승인'),
    ('도입 프로그램이', '학습관리시스템이'),
    ('도입 프로그램', '학습관리시스템'),
    ('고객 현장 교육', '현장 교육'),
    ('고객 현장에서', '사업장 현장에서'),
    ('동일 고객 건', '동일 부서 건'),
    ('고객사', '요청 부서'),
    # 사내 교육팀이 쓰는 말로 바꾼다. 남은 '고객'과 익명화 자국을 정리한다.
    ('활성화 코드 수령 후 고객과 배포 일정을', '활성화 코드 수령 후 IT본부와 배포 일정을'),
    ('무료 사용자 대상 기능 제공 방식 변경', '기본 라이선스 사용자 대상 기능 제공 방식 변경'),
    ('무료 사용자 참석자에게는', '기본 라이선스 참석자에게는'),
    ('요청번호 R-0000', '요청번호 SR-2418'),
]

hits = [0]


def swap(text):
    out = text
    for a, b in MAP:
        if a in out:
            hits[0] += out.count(a)
            out = out.replace(a, b)
    return out


def fix_para(p):
    """문단 안에서 run 단위로 바꾼다. 걸치면 첫 run에 몰아 넣는다."""
    if not any(a in p.text for a, _ in MAP):
        return
    for r in p.runs:
        new = swap(r.text)
        if new != r.text:
            r.text = new
    if any(a in p.text for a, _ in MAP) and p.runs:
        merged = swap(''.join(r.text for r in p.runs))
        p.runs[0].text = merged
        for r in p.runs[1:]:
            r.text = ''


def main():
    d = Document(DOC)
    for p in d.paragraphs:
        fix_para(p)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    fix_para(p)
    d.save(DOC)

    left = re.compile('고객사? [A-H]|도입 프로그램|고객|R-0000')
    rest = [p.text for p in Document(DOC).paragraphs if left.search(p.text)]
    print('치환 %d곳' % hits[0])
    if rest:
        print('남은 곳:', rest[:3])
        sys.exit(1)
    print('남은 대체어 없음')


if __name__ == '__main__':
    main()
